from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "backend-optimization-requirements-phase-2.md"
OUTPUT = ROOT / "German-Steels-Backend-Optimization-Requirements-Phase-2.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "C9D2DC"
WHITE = "FFFFFF"
BLACK = "000000"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM_DXA, start=CELL_SIDE_DXA,
                     bottom=CELL_TOP_BOTTOM_DXA, end=CELL_SIDE_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)


def paragraph_shading(paragraph, fill, left_border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if left_border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        p_bdr.append(left)


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "GS Bullet" not in styles:
        bullet = styles.add_style("GS Bullet", 1)
    else:
        bullet = styles["GS Bullet"]
    bullet.base_style = normal
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    if "GS Number" not in styles:
        number = styles.add_style("GS Number", 1)
    else:
        number = styles["GS Number"]
    number.base_style = normal
    number.font.name = "Calibri"
    number.font.size = Pt(11)
    number.paragraph_format.left_indent = Inches(0.5)
    number.paragraph_format.first_line_indent = Inches(-0.25)
    number.paragraph_format.space_after = Pt(8)
    number.paragraph_format.line_spacing = 1.167

    if "GS Code" not in styles:
        code = styles.add_style("GS Code", 1)
    else:
        code = styles["GS Code"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("1F2937")
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    if "GS Metadata" not in styles:
        metadata = styles.add_style("GS Metadata", 1)
    else:
        metadata = styles["GS Metadata"]
    metadata.base_style = normal
    metadata.font.name = "Calibri"
    metadata.font.size = Pt(10.5)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.0

    if "GS Table Text" not in styles:
        table_text = styles.add_style("GS Table Text", 1)
    else:
        table_text = styles["GS Table Text"]
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.08


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text, default_size=None):
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            if default_size:
                set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9 if default_size is None else min(default_size, 9), color=DARK_BLUE)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if default_size:
            set_run_font(run, size=default_size)


def add_header_footer(section):
    header = section.header
    header.distance = Inches(0.492)
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    left = p.add_run("GERMAN STEELS  |  BACKEND OPTIMIZATION")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("PHASE 2  •  v1.0")
    set_run_font(right, size=8.5, color=MUTED, bold=True)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), 2)

    footer = section.footer
    footer.distance = Inches(0.492)
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.line_spacing = 1.0
    left_run = fp.add_run("Internal implementation specification")
    set_run_font(left_run, size=8.5, color=MUTED)
    fp.add_run("\t")
    add_page_field(fp)
    for run in fp.runs[1:]:
        set_run_font(run, size=8.5, color=MUTED)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)


def add_title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("TECHNICAL IMPLEMENTATION SPECIFICATION")
    set_run_font(run, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("German Steels Backend\nOptimization Requirements")
    set_run_font(run, size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True
    run = subtitle.add_run("Phase 2 — actionable changes for scalable lists, exports, aggregation, and access control")
    set_run_font(run, size=13, color=MUTED)

    metadata = [
        ("To", "Backend Engineering"),
        ("From", "German Steels Web Engineering"),
        ("Version", "1.0"),
        ("Date", "12 August 2026"),
        ("Backend", "http://ec2-18-211-58-135.compute-1.amazonaws.com:8081"),
        ("Status", "Implementation requested"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph(style="GS Metadata")
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=BLACK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_markdown_table(doc, rows):
    headers = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)

    if len(headers) == 4:
        widths = [840, 2500, 4200, 1820]
    elif len(headers) == 3:
        widths = [1600, 3900, 3860]
    elif len(headers) == 2:
        widths = [2700, 6660]
    else:
        base = CONTENT_WIDTH_DXA // len(headers)
        widths = [base] * len(headers)
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)

    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.style = "GS Table Text"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index else WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, value, default_size=9)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])

    for row_values in body:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.style = "GS Table Text"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, default_size=9)

    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def build_document():
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet")
    add_header_footer(section)
    add_title_block(doc)

    # Skip the source title, subtitle, and metadata already rendered above.
    i = 10
    in_code = False
    code_lines = []
    active_number_num_id = None
    numbering_active = False

    while i < len(source_lines):
        line = source_lines[i]

        if line.startswith("```"):
            numbering_active = False
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph(style="GS Code")
                paragraph_shading(p, LIGHT_GRAY)
                code_run = p.add_run("\n".join(code_lines))
                set_run_font(code_run, name="Consolas", size=8.5, color="1F2937")
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("|"):
            numbering_active = False
            table_lines = []
            while i < len(source_lines) and source_lines[i].startswith("|"):
                table_lines.append(source_lines[i])
                i += 1
            parsed = []
            for row_index, table_line in enumerate(table_lines):
                values = [value.strip() for value in table_line.strip().strip("|").split("|")]
                if row_index == 1 and all(re.fullmatch(r":?-{3,}:?", value) for value in values):
                    continue
                parsed.append(values)
            add_markdown_table(doc, parsed)
            continue

        if line.startswith("> "):
            numbering_active = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.08)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.10
            paragraph_shading(p, CALLOUT, BLUE)
            add_inline(p, line[2:])
            i += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            numbering_active = False
            hashes, text = heading.groups()
            style_name = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[len(hashes)]
            p = doc.add_paragraph(style=style_name)
            add_inline(p, text)
            i += 1
            continue

        if line.startswith("- "):
            numbering_active = False
            p = doc.add_paragraph(style="GS Bullet")
            apply_numbering(p, bullet_num_id)
            add_inline(p, line[2:])
            i += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            if not numbering_active:
                active_number_num_id = add_numbering_definition(doc, "decimal")
                numbering_active = True
            p = doc.add_paragraph(style="GS Number")
            apply_numbering(p, active_number_num_id)
            add_inline(p, number_match.group(1))
            i += 1
            continue

        numbering_active = False
        p = doc.add_paragraph()
        keep_candidate = line.replace("`", "")
        if re.match(r"^(GET|POST|PUT|PATCH|DELETE)\s+/", keep_candidate) or line.rstrip().endswith(":"):
            p.paragraph_format.keep_with_next = True
        add_inline(p, line)
        i += 1

    doc.core_properties.title = "German Steels Backend Optimization Requirements - Phase 2"
    doc.core_properties.subject = "Backend implementation specification"
    doc.core_properties.author = "German Steels Web Engineering"
    doc.core_properties.keywords = "German Steels, backend, API, optimization, pagination, export"
    doc.core_properties.comments = "Prepared from an audit of the active German Steels web application."
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
